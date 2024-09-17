from docx.shared import Inches
from ebooklib import epub
from docx import Document

class DocumentCreator:
    @staticmethod
    def create_docx(text, image_files, output_file):
        """Tworzy plik DOCX z tekstem i obrazami"""
        doc = Document()
        doc.add_paragraph(text)
        for image in image_files:
            doc.add_picture(image, width=Inches(4))
        doc.save(output_file)

    @staticmethod
    def create_epub(text, image_files, output_file):
        """Tworzy plik EPUB z tekstem i obrazami"""
        book = epub.EpubBook()

        # Ustawienia metadanych książki
        book.set_identifier('id123456')
        book.set_title('Generated Book')
        book.set_language('en')

        # Dodanie rozdziału z tekstem
        chapter = epub.EpubHtml(title='Chapter 1', file_name='chap_01.xhtml', lang='en')
        chapter.content = f'<h1>Chapter 1</h1><p>{text}</p>'
        book.add_item(chapter)

        # Dodawanie obrazów jako osobnych rozdziałów
        for idx, image in enumerate(image_files):
            image_chapter = epub.EpubHtml(title=f'Image Chapter {idx+1}', file_name=f'image_chap_{idx+1}.xhtml')
            image_chapter.content = f'<img src="{image}" />'
            book.add_item(image_chapter)

        # Zdefiniuj style CSS (opcjonalnie)
        style = epub.EpubItem(uid="style_nav", file_name="style/style.css", media_type="text/css", content='body { font-family: Arial; }')
        book.add_item(style)

        # Dodaj niezbędne elementy do książki
        book.toc = (epub.Link('chap_01.xhtml', 'Chapter 1', 'chapter1'),
                    (epub.Section('Images'),
                     [epub.Link(f'image_chap_{idx+1}.xhtml', f'Image Chapter {idx+1}', f'imgchap{idx+1}') for idx in range(len(image_files))]))

        # Zdefiniowanie spine (kręgosłupa książki)
        book.spine = ['nav', chapter] + [epub.EpubNcx(), epub.EpubNav()]

        # Zapisz książkę EPUB
        epub.write_epub(output_file, book, {})

